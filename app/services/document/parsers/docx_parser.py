import logging
from docx import Document as DocxDocument
from .base_parser import BaseParser, ParseResult

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Word 文档解析器，使用 python-docx"""

    def parse(self, file_path: str) -> ParseResult:
        doc = DocxDocument(file_path)

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(' | '.join(row_text))

        content = '\n\n'.join(paragraphs)

        # 提取元数据
        core_props = doc.core_properties
        title = core_props.title or ''
        if not title:
            import os
            title = os.path.splitext(os.path.basename(file_path))[0]

        return ParseResult(
            content=content.strip(),
            title=title,
            metadata={
                'author': core_props.author or '',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'format': 'docx',
            }
        )

    def supported_extensions(self) -> list:
        return ['.docx', '.doc']
